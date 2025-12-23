"""
Document Loader Module for Construction RAG System
===================================================

This module handles loading various document formats commonly found in 
construction projects: PDFs, DOCX, TXT, and potentially HTML/XML.

Design Decisions:
-----------------
1. Factory Pattern: We use a factory to create appropriate loaders based on 
   file extension. This is extensible - adding new formats is trivial.

2. Lazy Loading: Documents are loaded on-demand with generators to handle 
   large document sets without memory issues.

3. Metadata Extraction: We extract rich metadata (filename, page numbers, 
   creation date) which is crucial for:
   - Citation in responses (grounding)
   - Filtering during retrieval
   - Debugging and traceability

ML Fundamentals Connection:
---------------------------
Document loading is the data preprocessing stage of our ML pipeline.
Just as image ML requires loading, resizing, and normalizing images,
RAG requires loading, cleaning, and structuring documents.

The quality of ingestion directly impacts model performance
(Garbage In, Garbage Out principle).
"""

import os
import hashlib
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Generator, Optional
from dataclasses import dataclass, field

from loguru import logger


@dataclass
class Document:
    """
    Represents a loaded document with content and metadata.
    
    This is our fundamental data structure that flows through the pipeline.
    Think of it like a DataFrame in pandas - a standardized container for data.
    
    Attributes:
        content: The raw text content of the document
        metadata: Rich metadata for traceability and filtering
        doc_id: Unique identifier (hash of content + source)
    """
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = ""
    
    def __post_init__(self):
        """Generate unique document ID after initialization."""
        if not self.doc_id:
            # Create deterministic ID from content + source
            id_source = f"{self.content[:500]}{self.metadata.get('source', '')}"
            self.doc_id = hashlib.md5(id_source.encode()).hexdigest()[:12]
    
    def __repr__(self) -> str:
        preview = self.content[:100].replace('\n', ' ')
        return f"Document(id={self.doc_id}, preview='{preview}...')"


class BaseDocumentLoader(ABC):
    """
    Abstract base class for document loaders.
    
    Design Pattern: Template Method
    - Defines the skeleton of document loading algorithm
    - Subclasses implement specific format handling
    """
    
    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        """Load document(s) from a file path."""
        pass
    
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        pass
    
    def _extract_base_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract common metadata from any file."""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            "source": str(path.absolute()),
            "filename": path.name,
            "file_extension": path.suffix.lower(),
            "file_size_bytes": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "loaded_at": datetime.now().isoformat(),
        }


class TextDocumentLoader(BaseDocumentLoader):
    """
    Loader for plain text files (.txt, .md, .csv).
    
    This is the simplest loader - directly reads file content.
    Used as baseline and for testing.
    """
    
    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding
    
    def supported_extensions(self) -> List[str]:
        return [".txt", ".md", ".text", ".csv"]
    
    def load(self, file_path: str) -> List[Document]:
        """Load a text file and return as Document."""
        logger.debug(f"Loading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding=self.encoding) as f:
                content = f.read()
            
            metadata = self._extract_base_metadata(file_path)
            metadata["loader_type"] = "text"
            metadata["char_count"] = len(content)
            metadata["line_count"] = content.count('\n') + 1
            
            logger.info(f"Loaded text file: {file_path} ({metadata['char_count']} chars)")
            
            return [Document(content=content, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise


class PDFDocumentLoader(BaseDocumentLoader):
    """
    Loader for PDF documents.
    
    PDFs are ubiquitous in construction (contracts, specs, drawings).
    We use pypdf for basic extraction.
    
    Trade-off Analysis:
    - pypdf: Lightweight, fast, good for text-heavy PDFs
    - pdfplumber: Better for tables and complex layouts
    - unstructured: Best quality but slower, heavier dependencies
    
    We default to pypdf for balance of speed and quality.
    """
    
    def supported_extensions(self) -> List[str]:
        return [".pdf"]
    
    def load(self, file_path: str, pages_per_doc: Optional[int] = None) -> List[Document]:
        """
        Load PDF and optionally split into multiple documents by page.
        
        Args:
            file_path: Path to PDF file
            pages_per_doc: If set, create separate Document per N pages.
                          None means entire PDF as one Document.
        """
        logger.debug(f"Loading PDF: {file_path}")
        
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.error("pypdf not installed. Run: pip install pypdf")
            raise
        
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            documents = []
            
            base_metadata = self._extract_base_metadata(file_path)
            base_metadata["loader_type"] = "pdf"
            base_metadata["total_pages"] = total_pages
            
            if pages_per_doc is None:
                # Load entire PDF as one document
                content_parts = []
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    content_parts.append(f"[Page {page_num + 1}]\n{text}")
                
                metadata = base_metadata.copy()
                metadata["page_range"] = f"1-{total_pages}"
                
                documents.append(Document(
                    content="\n\n".join(content_parts),
                    metadata=metadata
                ))
            else:
                # Split into multiple documents
                for start in range(0, total_pages, pages_per_doc):
                    end = min(start + pages_per_doc, total_pages)
                    content_parts = []
                    
                    for page_num in range(start, end):
                        text = reader.pages[page_num].extract_text() or ""
                        content_parts.append(f"[Page {page_num + 1}]\n{text}")
                    
                    metadata = base_metadata.copy()
                    metadata["page_range"] = f"{start + 1}-{end}"
                    metadata["is_partial"] = True
                    
                    documents.append(Document(
                        content="\n\n".join(content_parts),
                        metadata=metadata
                    ))
            
            logger.info(f"Loaded PDF: {file_path} ({total_pages} pages -> {len(documents)} documents)")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise


class DOCXDocumentLoader(BaseDocumentLoader):
    """
    Loader for Microsoft Word documents.
    
    Word documents are common for contracts, reports, and specifications
    in construction projects.
    """
    
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]
    
    def load(self, file_path: str) -> List[Document]:
        """Load DOCX file and extract text content."""
        logger.debug(f"Loading DOCX: {file_path}")
        
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_content.append(row_text)
                table_texts.append("\n".join(table_content))
            
            # Combine content
            content_parts = paragraphs + table_texts
            content = "\n\n".join(content_parts)
            
            metadata = self._extract_base_metadata(file_path)
            metadata["loader_type"] = "docx"
            metadata["paragraph_count"] = len(paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            # Try to extract core properties
            try:
                if doc.core_properties:
                    metadata["author"] = doc.core_properties.author or ""
                    metadata["title"] = doc.core_properties.title or ""
                    if doc.core_properties.created:
                        metadata["doc_created"] = doc.core_properties.created.isoformat()
            except:
                pass  # Core properties may not be available
            
            logger.info(f"Loaded DOCX: {file_path} ({len(paragraphs)} paragraphs, {len(doc.tables)} tables)")
            
            return [Document(content=content, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise


class DocumentLoaderFactory:
    """
    Factory for creating appropriate document loaders.
    
    Design Pattern: Factory Pattern
    - Centralizes loader creation logic
    - Makes adding new formats trivial
    - Client code doesn't need to know about specific loaders
    """
    
    def __init__(self):
        """Initialize factory with available loaders."""
        self._loaders = {}
        self._register_default_loaders()
    
    def _register_default_loaders(self):
        """Register built-in document loaders."""
        loaders = [
            TextDocumentLoader(),
            PDFDocumentLoader(),
            DOCXDocumentLoader(),
        ]
        
        for loader in loaders:
            for ext in loader.supported_extensions():
                self._loaders[ext] = loader
    
    def register_loader(self, extension: str, loader: BaseDocumentLoader):
        """Register a custom loader for a file extension."""
        self._loaders[extension.lower()] = loader
    
    def get_loader(self, file_path: str) -> BaseDocumentLoader:
        """Get appropriate loader for a file."""
        ext = Path(file_path).suffix.lower()
        
        if ext not in self._loaders:
            raise ValueError(f"No loader registered for extension: {ext}")
        
        return self._loaders[ext]
    
    def supported_extensions(self) -> List[str]:
        """Return all supported file extensions."""
        return list(self._loaders.keys())


class DirectoryLoader:
    """
    Load all supported documents from a directory.
    
    This is the main entry point for batch document loading.
    Recursively finds and loads all supported files.
    """
    
    def __init__(
        self,
        directory: str,
        recursive: bool = True,
        exclude_patterns: Optional[List[str]] = None
    ):
        self.directory = Path(directory)
        self.recursive = recursive
        self.exclude_patterns = exclude_patterns or []
        self.factory = DocumentLoaderFactory()
    
    def _should_exclude(self, path: Path) -> bool:
        """Check if file should be excluded."""
        path_str = str(path)
        for pattern in self.exclude_patterns:
            if pattern in path_str:
                return True
        return False
    
    def load(self) -> Generator[Document, None, None]:
        """
        Load all documents from directory.
        
        Uses a generator for memory-efficient processing of large directories.
        This is crucial for production systems with thousands of documents.
        """
        if not self.directory.exists():
            raise FileNotFoundError(f"Directory not found: {self.directory}")
        
        supported = set(self.factory.supported_extensions())
        
        # Find all files
        if self.recursive:
            files = self.directory.rglob("*")
        else:
            files = self.directory.glob("*")
        
        # Process each file
        for file_path in files:
            if not file_path.is_file():
                continue
            
            ext = file_path.suffix.lower()
            if ext not in supported:
                logger.debug(f"Skipping unsupported file: {file_path}")
                continue
            
            if self._should_exclude(file_path):
                logger.debug(f"Excluding file: {file_path}")
                continue
            
            try:
                loader = self.factory.get_loader(str(file_path))
                documents = loader.load(str(file_path))
                
                for doc in documents:
                    yield doc
                    
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                # Continue with other files instead of failing
                continue
    
    def load_all(self) -> List[Document]:
        """Load all documents and return as list."""
        return list(self.load())


# Convenience function for simple usage
def load_documents(
    path: str,
    recursive: bool = True
) -> List[Document]:
    """
    Simple function to load documents from file or directory.
    
    Args:
        path: Path to file or directory
        recursive: If directory, load recursively
    
    Returns:
        List of Document objects
    """
    path_obj = Path(path)
    
    if path_obj.is_file():
        factory = DocumentLoaderFactory()
        loader = factory.get_loader(path)
        return loader.load(path)
    
    elif path_obj.is_dir():
        loader = DirectoryLoader(path, recursive=recursive)
        return loader.load_all()
    
    else:
        raise FileNotFoundError(f"Path not found: {path}")


if __name__ == "__main__":
    # Example usage and testing
    import sys
    
    logging_format = "<green>{time:HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>"
    logger.remove()
    logger.add(sys.stderr, format=logging_format, level="DEBUG")
    
    # Test with sample documents
    sample_dir = Path(__file__).parent.parent.parent.parent / "data" / "sample_docs"
    
    if sample_dir.exists():
        print(f"\n📁 Loading documents from: {sample_dir}\n")
        
        documents = load_documents(str(sample_dir))
        
        print(f"✅ Loaded {len(documents)} documents:\n")
        for doc in documents:
            print(f"  📄 {doc.metadata.get('filename', 'Unknown')}")
            print(f"     ID: {doc.doc_id}")
            print(f"     Size: {len(doc.content):,} characters")
            print(f"     Preview: {doc.content[:100]}...")
            print()
    else:
        print(f"Sample directory not found: {sample_dir}")
